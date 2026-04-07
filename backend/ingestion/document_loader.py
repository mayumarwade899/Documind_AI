import os
import hashlib
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Optional
from datetime import datetime

import pdfplumber
from pypdf import PdfReader 
from docx import Document as DocxDocument
from config.logging_config import get_logger

logger = get_logger(__name__)

@dataclass
class DocumentPage:

    content: str
    page_number: int
    source_file: str
    file_type: str
    document_id: str
    total_page: int
    metadata: dict = field(default_factory=dict)

@dataclass
class LoadedDocument:
    document_id: str
    filename: str
    file_type: str
    file_path: str
    total_pages: int
    pages: List[DocumentPage]
    loaded_at: str
    file_size_bytes: int

def _generate_document_id(file_path: str) -> str:
    h = hashlib.md5()
    with open(file_path, "rb") as f:
        while True:
            chunk = f.read(65536)
            if not chunk:
                break
            h.update(chunk)
    return h.hexdigest()

def _clean_text(text: str) -> str:
    if not text:
        return ""

    text = text.replace("\x00", "")

    lines = text.splitlines()
    cleaned = []
    prev_blank = False

    for line in lines:
        stripped = line.strip()
        if stripped == "":
            if not prev_blank:
                cleaned.append("")
            prev_blank = True
        else:
            cleaned.append(stripped)
            prev_blank = False

    return "\n".join(cleaned).strip()

def _load_pdf(file_path: str, document_id: str) -> List[DocumentPage]:
    filename = Path(file_path).name
    pages = []

    pypdf_reader = None  # Lazy-loaded only when pdfplumber fails

    try:
        plumber_doc = pdfplumber.open(file_path)
        total_pages = len(plumber_doc.pages)

        logger.info(
            "pdf_loading_started",
            filename = filename,
            total_pages = total_pages
        )

        for page_num in range(total_pages):
            raw_text = ""
            try:
                plumber_page = plumber_doc.pages[page_num]
                raw_text = plumber_page.extract_text() or ""
            except Exception as e:
                logger.warning(
                    "pdfplumber_page_failed",
                    filename = filename,
                    page = page_num + 1,
                    error = str(e)
                )
            
            if not raw_text.strip():
                try:
                    if pypdf_reader is None:
                        pypdf_reader = PdfReader(file_path)
                    pypdf_page = pypdf_reader.pages[page_num]
                    raw_text = pypdf_page.extract_text() or ""
                    if raw_text.strip():
                        logger.debug(
                            "pypdf_fallback_success",
                            filename = filename,
                            page = page_num + 1
                        )
                except Exception as e:
                    logger.warning(
                        "pypdf_fallback_failed",
                        filename = filename,
                        page = page_num + 1,
                        error = str(e)
                    )
            
            clean = _clean_text(raw_text)

            if not clean:
                logger.debug(
                    "pdf_page_skipped_empty",
                    filename = filename,
                    page = page_num + 1
                )
                continue

            pages.append(
                DocumentPage(
                    content = clean,
                    page_number = page_num + 1,
                    source_file = filename,
                    file_type = "pdf",
                    document_id = document_id,
                    total_page = total_pages,
                    metadata = {"parser": "pdfplumber" if raw_text else "pypdf"}
                )
            )
            
        plumber_doc.close()
        logger.info(
            "pdf_loading_complete",
            filename=filename,
            pages_extracted=len(pages)
        )

    except Exception as e:
        logger.error("pdf_loading_failed", filename = filename, error = str(e))
        raise

    return pages

def _load_txt(file_path: str, document_id: str) -> List[DocumentPage]:
    filename = Path(file_path).name

    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            raw_text = f.read()

        clean = _clean_text(raw_text)

        if not clean:
            logger.warning("text_file_empty", filename = filename)
            return []

        logger.info("txt_loading_complete", filename = filename)

        return [DocumentPage(
            content = clean,
            page_number = 1,
            source_file = filename,
            file_type = "txt",
            document_id = document_id,
            total_pages = 1,
            metadata = {}
        )]
    
    except Exception as e:
        logger.error("txt_loading_failed", filename=filename, error=str(e))
        raise

def _load_docx(file_path: str, document_id: str) -> List[DocumentPage]:
    filename = Path(file_path).name

    try:
        doc = DocxDocument(file_path)

        paragraphs = [
            p.text.strip()
            for p in doc.paragraphs
            if p.text.strip()
        ]

        if not paragraphs:
            logger.warning("docx_file_empty", filename = filename)
            return []

        PAGE_SIZE = 50
        pages = []
        total_pages = (len(paragraphs) + PAGE_SIZE - 1)

        for i in range(0, len(paragraphs), PAGE_SIZE):
            content = _clean_text("/n".join(paragraphs[i:i + PAGE_SIZE]))
            if content:
                pages.append(DocumentPage(
                    content = content,
                    page_number = (i // PAGE_SIZE) + 1,
                    source_file = filename,
                    file_type = "docx",
                    document_id=document_id,
                    total_pages=total_pages,
                    metadata={}
                ))
        
        logger.info(
            "docx_loading_complete",
            filename = filename,
            pages_extracted = len(pages)
        )
        return pages

    except Exception as e:
        logger.error("docx_loading_failed", filename=filename, error=str(e))
        raise

class DocumentLoader:
    SUPPORTED_EXTENSIONS = {
        ".pdf": _load_pdf,
        ".txt": _load_txt,
        ".docx": _load_docx,
    }

    def load(self, file_path: str) -> LoadedDocument:
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type '{ext}'. "
                f"Supported: {list(self.SUPPORTED_EXTENSIONS.keys())}"
            )
            
        logger.info("document_load_started", file_path = str(path))

        document_id = _generate_document_id(file_path)
        file_size = os.path.getsize(file_path)
        pages = self.SUPPORTED_EXTENSIONS[ext](file_path, document_id)

        doc = LoadedDocument(
            document_id = document_id,
            filename = path.name,
            file_type = ext.lstrip("."),
            file_path = str(path.absolute()),
            total_pages = len(pages),
            pages = pages,
            loaded_at = datetime.utcnow().isoformat(),
            file_size_bytes = file_size
        )

        logger.info(
            "document_load_complete",
            filename = path.name,
            document_id = document_id,
            pages = len(pages),
            size_kb = round(file_size / 1024, 2)
        )

        return doc
    
    def load_directory(self, dir_path: str) -> List[LoadedDocument]:
        directory = Path(dir_path)

        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {dir_path}")

        files = [
            f for f in directory.iterdir()
            if f.suffix.lower() in self.SUPPORTED_EXTENSIONS
        ]

        if not files:
            logger.warning("No documents found", dir_path = dir_path)
            return []

        logger.info(
            "directory_load_started",
            dir_path = dir_path,
            file_count = len(files)
        )

        documents = []

        for file_path in files:
            try:
                documents.append(self.load(str(file_path)))
            except Exception as e:
                logger.error(
                    "documents_skipped",
                    filename = file_path.name,
                    error = str(e)
                )

        logger.info(
            "directory_load_complete",
            loaded = len(documents),
            skipped = len(files) - len(documents)
        )

        return documents
